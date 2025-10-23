"""In-memory reporting bundle helpers for Streamlit downloads."""

from __future__ import annotations

import io
import logging
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional, Tuple, TYPE_CHECKING, Set

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from openpyxl.utils import get_column_letter

from .report_generator import ReportGenerator
from ..models.results import EngineResults
from ..visualization import plot_shock_group_summary

if TYPE_CHECKING:
    from ..engine import DiscountConfig

LOGGER = logging.getLogger(__name__)


@dataclass
class InMemoryReportBundle:
    """Container for downloadable report artifacts."""

    zip_bytes: bytes
    zip_name: str
    excel_name: str
    word_name: str
    manifest: Dict[str, str]
    created_at: datetime


class InMemoryReportBuilder:
    """Construct Excel + Word deliverables without touching the filesystem."""

    FIGURE_LABELS: Dict[str, str] = {
        "rate_paths": "Rate Paths",
        "shock_magnitude": "Shock Magnitude",
        "tenor_comparison": "Tenor Comparison",
        "pv_delta": "PV Delta Comparison",
        "rate_spaghetti": "Simulated Rate Paths",
        "rate_fan": "Rate Confidence Fan",
        "pv_distribution": "Portfolio PV Distribution",
        "percentiles": "Percentile Ladder",
        "dashboard": "Monte Carlo Dashboard",
    }

    def __init__(self, *, base_title: str = "Deposit Analysis", timestamp: Optional[str] = None) -> None:
        self.generated_at = datetime.utcnow()
        self._safe_title = self._sanitize_title(base_title)
        self._slug_title = self._slugify(self._safe_title)
        self.timestamp = timestamp or self.generated_at.strftime("%Y%m%d_%H%M%S")

    # ------------------------------------------------------------------ helpers
    @staticmethod
    def _sanitize_title(title: str) -> str:
        cleaned = re.sub(r'[<>:"/\\|?*]+', "", title).strip()
        cleaned = re.sub(r"\s+", " ", cleaned)
        return cleaned or "Deposit Analysis"

    @staticmethod
    def _slugify(text: str) -> str:
        slug = re.sub(r"[^A-Za-z0-9]+", "_", text.strip())
        return slug.strip("_") or "report"

    @staticmethod
    def _safe_sheet_name(name: str, used: Set[str]) -> str:
        sanitized = re.sub(r'[\[\]\\/*?:]+', "-", name).strip()
        sanitized = re.sub(r"\s+", " ", sanitized)
        sanitized = sanitized or "Sheet"
        base = sanitized[:31]
        candidate = base
        counter = 1
        while candidate in used:
            suffix = f" ({counter})"
            candidate = f"{base[: 31 - len(suffix)]}{suffix}"
            counter += 1
        used.add(candidate)
        return candidate

    @staticmethod
    def _ensure_unique_path(path: str, used: Set[str]) -> str:
        if path not in used:
            used.add(path)
            return path
        stem, dot, ext = path.rpartition(".")
        stem = stem or path
        ext = ext or ""
        counter = 1
        while True:
            candidate = f"{stem}_{counter}"
            if ext:
                candidate = f"{candidate}.{ext}"
            if candidate not in used:
                used.add(candidate)
                return candidate
            counter += 1

    @staticmethod
    def _safe_component_name(text: str, *, extension: str = "png", folder: Optional[str] = None) -> str:
        cleaned = re.sub(r"[<>:\"/\\|?*\s]+", "_", text.strip())
        cleaned = cleaned.strip("_") or "artifact"
        filename = f"{cleaned}.{extension}"
        if folder:
            return f"{folder.rstrip('/')}/{filename}"
        return filename

    @staticmethod
    def _normalize_scalar(value: Any) -> Any:
        if isinstance(value, (np.integer,)):
            return int(value)
        if isinstance(value, (np.floating,)):
            return float(value)
        if isinstance(value, (datetime,)):
            return value.isoformat()
        if isinstance(value, pd.Timestamp):
            return value.isoformat()
        if isinstance(value, pd.Timedelta):
            return value.isoformat()
        if isinstance(value, Path):
            return str(value)
        return value

    @classmethod
    def _flatten_mapping(cls, data: Any, prefix: Tuple[str, ...] = ()) -> List[Tuple[str, Any]]:
        rows: List[Tuple[str, Any]] = []
        if isinstance(data, Mapping):
            for key, value in sorted(data.items(), key=lambda item: str(item[0])):
                rows.extend(cls._flatten_mapping(value, prefix + (str(key),)))
            return rows
        if isinstance(data, (list, tuple)):
            for idx, value in enumerate(data):
                rows.extend(cls._flatten_mapping(value, prefix + (f"[{idx}]",)))
            return rows
        path = " / ".join(prefix) if prefix else "value"
        rows.append((path, cls._normalize_scalar(data)))
        return rows

    @classmethod
    def _mapping_to_frame(
        cls,
        mapping: Mapping[str, Any],
        *,
        columns: Tuple[str, str] = ("Field", "Value"),
    ) -> pd.DataFrame:
        if not mapping:
            return pd.DataFrame(columns=columns)
        flattened = cls._flatten_mapping(mapping)
        if not flattened:
            return pd.DataFrame(columns=columns)
        frame = pd.DataFrame(flattened, columns=columns)
        return frame

    @staticmethod
    def _scenario_label(scenario_id: str, metadata: Optional[Mapping[str, Any]]) -> str:
        if metadata:
            label = metadata.get("description")
            if label:
                return str(label)
        return scenario_id.replace("_", " ").title()

    @staticmethod
    def _figure_to_png(fig) -> bytes:
        buffer = io.BytesIO()
        fig.savefig(buffer, format="png", dpi=300, bbox_inches="tight")
        plt.close(fig)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def _format_value(column: str, value: Any) -> str:
        if value is None:
            return "-"
        if isinstance(value, float) and np.isnan(value):
            return "-"
        name = column.lower()
        if isinstance(value, (np.integer, int)):
            numeric = float(value)
            if any(token in name for token in ("bps", "basis")):
                return f"{numeric:.0f} bps"
            if any(token in name for token in ("year", "years")):
                return f"{numeric:,.2f} yrs"
            return f"{int(value):,}"
        if isinstance(value, (np.floating, float)):
            numeric = float(value)
            if any(token in name for token in ("bps", "basis")):
                return f"{numeric * 10000:.0f} bps"
            if any(token in name for token in ("beta", "ratio", "pct", "percent", "decay", "rate", "shock")) and abs(numeric) <= 1.5:
                return f"{numeric * 100:.2f}%"
            if any(token in name for token in ("year", "years")):
                return f"{numeric:,.2f} yrs"
            if any(token in name for token in ("month", "months")):
                return f"{numeric:,.0f} months"
            if any(token in name for token in ("present_value", "pv", "balance", "amount", "notional", "cash", "value", "delta")):
                return f"${numeric:,.2f}"
            return f"{numeric:,.4f}" if abs(numeric) < 1 else f"{numeric:,.2f}"
        if isinstance(value, (datetime, pd.Timestamp)):
            return value.isoformat()
        if isinstance(value, str):
            return value
        if isinstance(value, (list, tuple, set)):
            seq = list(value)
            if not seq:
                return "[]"
            if all(isinstance(v, (int, float, np.integer, np.floating)) for v in seq):
                return f"{len(seq)} values (min {min(seq):.2f}, max {max(seq):.2f})"
            return f"{len(seq)} values"
        return str(value)

    @classmethod
    def _add_dataframe_table(
        cls,
        document: Document,
        dataframe: pd.DataFrame,
        *,
        caption: Optional[str] = None,
        max_rows: int = 500,
    ) -> None:
        if caption:
            document.add_paragraph(caption, style="Caption")
        if dataframe is None or dataframe.empty:
            document.add_paragraph("No data available.")
            return
        display_df = dataframe.head(max_rows)
        table = document.add_table(rows=1, cols=len(display_df.columns))
        try:
            table.style = "Light List Accent 1"
        except Exception:  # pragma: no cover - style fallback
            pass
        header_cells = table.rows[0].cells
        for idx, column in enumerate(display_df.columns):
            header_cells[idx].text = str(column)
        for _, row in display_df.iterrows():
            cells = table.add_row().cells
            for idx, column in enumerate(display_df.columns):
                context = str(column)
                if column.lower() == "value" and "Field" in display_df.columns:
                    context = str(row.get("Field", column))
                cells[idx].text = cls._format_value(context, row.iloc[idx])
        document.add_paragraph("")

    def _excel_filename(self) -> str:
        return f"{self._safe_title} Data {self.timestamp}.xlsx"

    def _word_filename(self) -> str:
        return f"{self._safe_title} Report {self.timestamp}.docx"

    def _zip_filename(self) -> str:
        return f"{self._slug_title}_{self.timestamp}.zip"

    def _build_excel_workbook(self, sheet_frames: Dict[str, pd.DataFrame]) -> bytes:
        buffer = io.BytesIO()
        used_names: Set[str] = set()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            for name, frame in sheet_frames.items():
                if frame is None or frame.empty:
                    continue
                safe_name = self._safe_sheet_name(name, used_names)
                limited_frame = frame
                if len(limited_frame) > 1_048_576:
                    limited_frame = limited_frame.iloc[:1_048_576].copy()
                    LOGGER.warning(
                        "Sheet '%s' truncated to the first 1,048,576 rows for Excel compatibility.",
                        safe_name,
                    )
                limited_frame.to_excel(writer, sheet_name=safe_name, index=False)
                worksheet = writer.sheets[safe_name]
                row_count, col_count = limited_frame.shape
                if col_count > 0:
                    worksheet.freeze_panes = worksheet["A2"]
                    last_row = row_count + 1 if row_count > 0 else 1
                    last_col_letter = get_column_letter(col_count)
                    worksheet.auto_filter.ref = f"A1:{last_col_letter}{last_row}"
                    for col_idx, column in enumerate(limited_frame.columns, start=1):
                        sample = limited_frame[column].head(200)
                        max_len = max([len(str(column))] + [len(str(item)) for item in sample])
                        column_letter = get_column_letter(col_idx)
                        worksheet.column_dimensions[column_letter].width = min(max_len + 2, 60)
        buffer.seek(0)
        return buffer.getvalue()

    def _build_word_report(
        self,
        summary_df: pd.DataFrame,
        parameter_df: pd.DataFrame,
        validation_df: pd.DataFrame,
        analysis_df: pd.DataFrame,
        scenario_details: Dict[str, Dict[str, Any]],
        scenario_images: Dict[str, List[Tuple[str, str, bytes]]],
        group_image: Optional[Tuple[str, str, bytes]],
        discount_config: Optional["DiscountConfig"],
        discount_curve_df: Optional[pd.DataFrame],
        discount_meta_df: Optional[pd.DataFrame],
    ) -> bytes:
        document = Document()
        document.add_heading(self._safe_title, level=0)
        document.add_paragraph(f"Generated on {self.generated_at:%Y-%m-%d %H:%M:%S} UTC")
        document.add_paragraph("")

        if summary_df is not None and not summary_df.empty:
            document.add_heading("Present Value Summary", level=1)
            self._add_dataframe_table(document, summary_df)

        if parameter_df is not None and not parameter_df.empty:
            document.add_heading("Key Parameters", level=1)
            self._add_dataframe_table(document, parameter_df)

        if validation_df is not None and not validation_df.empty:
            document.add_heading("Validation Checks", level=1)
            self._add_dataframe_table(document, validation_df)

        if analysis_df is not None and not analysis_df.empty:
            document.add_heading("Analysis Metadata", level=1)
            self._add_dataframe_table(document, analysis_df)

        if discount_config is not None:
            document.add_heading("Discount Curve Configuration", level=1)
            document.add_paragraph(f"Method: {discount_config.method}")
            document.add_paragraph(f"Source: {discount_config.source}")
            if discount_meta_df is not None and not discount_meta_df.empty:
                self._add_dataframe_table(document, discount_meta_df)
            if discount_curve_df is not None and not discount_curve_df.empty:
                curve_display = discount_curve_df.copy()
                curve_display.columns = ["Tenor (months)", "Annual Rate"]
                curve_display["Annual Rate"] = curve_display["Annual Rate"].apply(
                    lambda v: f"{float(v) * 100:.3f}%"
                )
                self._add_dataframe_table(document, curve_display, caption="Discount Curve Points", max_rows=200)

        if scenario_details:
            document.add_heading("Scenario Highlights", level=1)
        for scenario_id, details in scenario_details.items():
            label = details["label"]
            document.add_heading(label, level=2)
            document.add_paragraph(f"Method: {details['method']}")
            document.add_paragraph(f"Present Value: {self._format_value('present_value', details['present_value'])}")
            if details.get("vs_base") is not None:
                delta_str = self._format_value("vs_base", details["vs_base"])
                pct = details.get("vs_base_pct")
                pct_str = f" ({self._format_value('vs_base_pct', pct)})" if pct is not None else ""
                document.add_paragraph(f"Δ vs Base: {delta_str}{pct_str}")
            metadata_frame = details.get("metadata_frame")
            if metadata_frame is not None and not metadata_frame.empty:
                filtered_metadata = metadata_frame.copy()
                if "Field" in filtered_metadata.columns:
                    filtered_metadata["Field"] = filtered_metadata["Field"].astype(str)
                    fields = filtered_metadata["Field"].str.lower()
                    exclude_mask = fields.str.contains(r"\[\d+\]") | fields.str.startswith("segments /")
                    filtered_metadata = filtered_metadata[~exclude_mask]
                    highlight_tokens = (
                        "method",
                        "description",
                        "shock",
                        "abs_max_bps",
                        "mean_bps",
                        "basis",
                        "beta",
                        "profile",
                    )
                    highlight_mask = filtered_metadata["Field"].str.lower().apply(
                        lambda value: any(token in value for token in highlight_tokens)
                    )
                    if highlight_mask.any():
                        filtered_metadata = filtered_metadata[highlight_mask]
                    if not filtered_metadata.empty:
                        filtered_metadata["Field"] = (
                            filtered_metadata["Field"]
                            .str.replace(r"\s*/\s*", " / ", regex=True)
                            .str.replace(r"\[\d+\]", "", regex=True)
                            .str.replace("_", " ")
                            .str.title()
                        )
                        filtered_metadata = filtered_metadata.head(40)
                if filtered_metadata is not None and not filtered_metadata.empty:
                    self._add_dataframe_table(
                        document,
                        filtered_metadata,
                        caption="Scenario Metadata",
                        max_rows=40,
                    )
            for image_label, _, image_bytes in scenario_images.get(scenario_id, []):
                document.add_paragraph(image_label, style="Heading 3")
                document.add_picture(io.BytesIO(image_bytes), width=Inches(6.0))
                document.add_paragraph("")

        if group_image is not None:
            label, _, image_bytes = group_image
            document.add_heading(label, level=2)
            document.add_picture(io.BytesIO(image_bytes), width=Inches(6.0))
            document.add_paragraph("")

        document.add_paragraph(
            "Remember to save the Excel workbook and this document locally to retain the full analysis output.",
            style="Intense Quote",
        )

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ------------------------------------------------------------------ public
    def build(
        self,
        results: EngineResults,
        *,
        discount_config: Optional["DiscountConfig"] = None,
        analysis_metadata: Optional[Mapping[str, Any]] = None,
        export_cashflows: bool = False,
        cashflow_mode: str = "sample",
        cashflow_sample_size: int = 20,
        cashflow_random_state: Optional[int] = 42,
    ) -> InMemoryReportBundle:
        sheet_frames: Dict[str, pd.DataFrame] = {}

        def add_sheet(name: str, frame: pd.DataFrame) -> None:
            if frame is None or frame.empty:
                return
            candidate = name
            suffix = 1
            while candidate in sheet_frames:
                suffix += 1
                candidate = f"{name} ({suffix})"
            sheet_frames[candidate] = frame

        summary_df = results.summary_frame()
        if not summary_df.empty:
            add_sheet("Summary - PV Comparison", summary_df)

        parameter_df = self._mapping_to_frame(results.parameter_summary or {})
        if not parameter_df.empty:
            add_sheet("Summary - Parameters", parameter_df)

        validation_df = self._mapping_to_frame(results.validation_summary or {})
        if not validation_df.empty:
            add_sheet("Summary - Validation Checks", validation_df)

        analysis_df = self._mapping_to_frame(analysis_metadata or {})
        if not analysis_df.empty:
            add_sheet("Summary - Analysis Metadata", analysis_df)

        discount_curve_df: Optional[pd.DataFrame] = None
        discount_meta_df: Optional[pd.DataFrame] = None
        if discount_config is not None:
            curve_map = discount_config.curve.to_dict()
            discount_curve_df = pd.DataFrame(
                {
                    "tenor_months": list(curve_map.keys()),
                    "annual_rate": list(curve_map.values()),
                }
            ).sort_values("tenor_months")
            add_sheet("Summary - Discount Curve", discount_curve_df)
            if discount_config.metadata:
                discount_meta_df = self._mapping_to_frame(discount_config.metadata)
                if not discount_meta_df.empty:
                    add_sheet("Summary - Discount Metadata", discount_meta_df)

        summary_lookup: Dict[str, Dict[str, Any]] = {}
        if not summary_df.empty:
            summary_lookup = summary_df.set_index("scenario_id").to_dict("index")

        scenario_details: Dict[str, Dict[str, Any]] = {}
        scenario_summaries: List[Dict[str, Any]] = []
        scenario_images: Dict[str, List[Tuple[str, str, bytes]]] = {}
        figure_files: Dict[str, bytes] = {}
        figure_manifest: Dict[str, str] = {}
        used_paths: Set[str] = set()
        shock_ids: List[str] = []

        for scenario_id, scenario in results.scenario_results.items():
            metadata = scenario.metadata or {}
            method = metadata.get("method", "base")
            label = self._scenario_label(scenario_id, metadata)
            summary_row = summary_lookup.get(scenario_id, {})
            present_value = float(scenario.present_value)
            vs_base = summary_row.get("vs_base")
            vs_base_pct = summary_row.get("vs_base_pct")
            metadata_frame = self._mapping_to_frame(metadata) if metadata else pd.DataFrame(columns=["Field", "Value"])

            scenario_details[scenario_id] = {
                "label": label,
                "method": method,
                "present_value": present_value,
                "vs_base": float(vs_base) if vs_base is not None else None,
                "vs_base_pct": float(vs_base_pct) if vs_base_pct is not None else None,
                "metadata_frame": metadata_frame,
            }

            scenario_summaries.append(
                {
                    "scenario_id": scenario_id,
                    "label": label,
                    "method": method,
                    "present_value": present_value,
                    "vs_base": scenario_details[scenario_id]["vs_base"],
                    "vs_base_pct": scenario_details[scenario_id]["vs_base_pct"],
                }
            )

            if not scenario.account_level_pv.empty:
                add_sheet(f"{label} - Account PV", scenario.account_level_pv)

            if export_cashflows:
                cashflows_df = scenario.cashflows
                if cashflow_mode != "full":
                    cashflows_df = ReportGenerator.sample_cashflows(
                        cashflows_df,
                        sample_size=cashflow_sample_size,
                        random_state=cashflow_random_state,
                    )
                if not cashflows_df.empty:
                    add_sheet(f"{label} - Cashflows", cashflows_df)

            extra_tables = scenario.extra_tables or {}
            for key, table in extra_tables.items():
                if table is None or table.empty:
                    continue
                table_label = key.replace("_", " ").title()
                add_sheet(f"{label} - {table_label}", table)

            if not metadata_frame.empty:
                add_sheet(f"{label} - Metadata", metadata_frame)

            if method not in {None, "base", "monte_carlo"}:
                shock_ids.append(scenario_id)
                figures = ReportGenerator._build_shock_figures(results, scenario_id)
                for key, fig in figures.items():
                    base_label = self.FIGURE_LABELS.get(key, key.replace("_", " ").title())
                    figure_name = f"{label} - {base_label}"
                    path = self._ensure_unique_path(
                        self._safe_component_name(figure_name, folder="figures"),
                        used_paths,
                    )
                    png_bytes = self._figure_to_png(fig)
                    figure_files[path] = png_bytes
                    figure_manifest[path] = f"{base_label} ({label})"
                    scenario_images.setdefault(scenario_id, []).append((base_label, path, png_bytes))
            elif method == "monte_carlo":
                figures = ReportGenerator._build_monte_carlo_figures(results, scenario_id=scenario_id, prefix=scenario_id)
                for key, fig in figures.items():
                    cleaned_key = key.replace(f"{scenario_id}_", "")
                    base_label = self.FIGURE_LABELS.get(cleaned_key, cleaned_key.replace("_", " ").title())
                    figure_name = f"{label} - {base_label}"
                    path = self._ensure_unique_path(
                        self._safe_component_name(figure_name, folder="figures"),
                        used_paths,
                    )
                    png_bytes = self._figure_to_png(fig)
                    figure_files[path] = png_bytes
                    figure_manifest[path] = f"{base_label} ({label})"
                    scenario_images.setdefault(scenario_id, []).append((base_label, path, png_bytes))

        if scenario_summaries:
            scenario_overview = pd.DataFrame(scenario_summaries)
            add_sheet("Summary - Scenario Overview", scenario_overview)

        group_image: Optional[Tuple[str, str, bytes]] = None
        if shock_ids:
            try:
                group_fig = plot_shock_group_summary(
                    results,
                    scenario_ids=shock_ids,
                    title="Deterministic Scenario Comparison",
                )
            except Exception as exc:  # pragma: no cover - visualization optional
                LOGGER.warning("Unable to build shock group summary: %s", exc)
                group_fig = None
            if group_fig is not None:
                label_text = "Deterministic Scenario Comparison"
                path = self._ensure_unique_path(
                    self._safe_component_name(label_text, folder="figures"),
                    used_paths,
                )
                png_bytes = self._figure_to_png(group_fig)
                figure_files[path] = png_bytes
                figure_manifest[path] = label_text
                group_image = (label_text, path, png_bytes)

        excel_bytes = self._build_excel_workbook(sheet_frames)
        word_bytes = self._build_word_report(
            summary_df,
            parameter_df,
            validation_df,
            analysis_df,
            scenario_details,
            scenario_images,
            group_image,
            discount_config,
            discount_curve_df,
            discount_meta_df,
        )

        excel_name = self._excel_filename()
        word_name = self._word_filename()
        zip_name = self._zip_filename()

        manifest: Dict[str, str] = {
            excel_name: "Excel workbook containing all tabular outputs.",
            word_name: "Narrative report with embedded visuals.",
        }
        manifest.update(figure_manifest)

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(excel_name, excel_bytes)
            zf.writestr(word_name, word_bytes)
            for path, payload in figure_files.items():
                zf.writestr(path, payload)

        zip_buffer.seek(0)
        return InMemoryReportBundle(
            zip_bytes=zip_buffer.getvalue(),
            zip_name=zip_name,
            excel_name=excel_name,
            word_name=word_name,
            manifest=manifest,
            created_at=self.generated_at,
        )
